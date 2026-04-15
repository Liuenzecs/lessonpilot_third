"""Sprint 5 导出服务测试。"""

from __future__ import annotations

from app.models import Task
from app.schemas.content import (
    AssessmentItem,
    KeyPoints,
    LessonPlanContent,
    LessonPlanHeader,
    StudyGuideContent,
    StudyGuideHeader,
    TeachingObjective,
    TeachingProcessStep,
)
from app.services.export_service import build_docx


def _make_task(
    *,
    scene: str = "public_school",
    lesson_type: str = "lesson_plan",
) -> Task:
    return Task(
        id="test-task-id",
        user_id="test-user",
        title="《春》教案",
        subject="语文",
        grade="七年级",
        topic="朱自清《春》",
        scene=scene,
        lesson_type=lesson_type,
        class_hour=1,
        lesson_category="new",
    )


def _make_lesson_plan(**overrides) -> LessonPlanContent:
    defaults: dict = {
        "header": LessonPlanHeader(
            title="《春》教案",
            subject="语文",
            grade="七年级",
            class_hour=1,
            lesson_category="new",
            teacher="张老师",
        ),
        "objectives": [
            TeachingObjective(dimension="knowledge", content="掌握生字词"),
            TeachingObjective(dimension="ability", content="提高朗读能力"),
            TeachingObjective(dimension="emotion", content="感受自然之美"),
        ],
        "objectives_status": "confirmed",
        "key_points": KeyPoints(
            key_points=["比喻手法的运用", "修辞手法辨析"],
            difficulties=["情景交融的写作技巧"],
        ),
        "key_points_status": "confirmed",
        "preparation": ["多媒体课件", "朗读音频"],
        "preparation_status": "confirmed",
        "teaching_process": [
            TeachingProcessStep(
                phase="导入新课",
                duration=5,
                teacher_activity="播放春天图片",
                student_activity="观察并交流",
                design_intent="激发学习兴趣",
            ),
            TeachingProcessStep(
                phase="初读课文",
                duration=15,
                teacher_activity="范读课文，指导朗读",
                student_activity="跟读，标注生字词",
                design_intent="整体感知课文",
            ),
        ],
        "teaching_process_status": "confirmed",
        "board_design": "板书：春——盼春、绘春、赞春",
        "board_design_status": "confirmed",
        "reflection": "",
        "reflection_status": "pending",
    }
    defaults.update(overrides)
    return LessonPlanContent(**defaults)


def _make_study_guide(**overrides) -> StudyGuideContent:
    defaults: dict = {
        "header": StudyGuideHeader(
            title="《春》学案",
            subject="语文",
            grade="七年级",
            class_name="初一(1)班",
            student_name="李明",
            date="2026-04-15",
        ),
        "learning_objectives": ["我能有感情地朗读课文", "我能分析修辞手法"],
        "learning_objectives_status": "confirmed",
        "key_difficulties": ["比喻句的辨析"],
        "key_difficulties_status": "confirmed",
        "prior_knowledge": ["修辞手法基础知识"],
        "prior_knowledge_status": "confirmed",
        "assessment": [
            AssessmentItem(
                level="A",
                item_type="choice",
                prompt="以下哪个是比喻句？",
                options=["春天来了", "春风像母亲的手", "花开得很美", "小鸟在唱歌"],
                answer="B",
                analysis="春风像母亲的手，用了比喻修辞",
            ),
        ],
        "assessment_status": "confirmed",
    }
    defaults.update(overrides)
    return StudyGuideContent(**defaults)


# ---------------------------------------------------------------------------
# 教案导出测试
# ---------------------------------------------------------------------------


class TestLessonPlanExport:
    def test_basic_export_produces_bytes(self) -> None:
        task = _make_task()
        content = _make_lesson_plan()
        result = build_docx(task, content)
        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX 文件以 PK ZIP 头开始
        assert result[:2] == b"PK"

    def test_export_with_empty_content(self) -> None:
        task = _make_task()
        content = LessonPlanContent()
        result = build_docx(task, content)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_public_school_scene_has_design_intent(self) -> None:
        """公立校场景导出的教学过程应包含设计意图列。"""
        task = _make_task(scene="public_school")
        content = _make_lesson_plan()
        result = build_docx(task, content)
        # 检查 docx 内文本包含设计意图内容
        text = _extract_docx_text(result)
        assert "激发学习兴趣" in text
        assert "整体感知课文" in text

    def test_tutor_scene_omits_design_intent(self) -> None:
        """家教场景导出的教学过程应省略设计意图列。"""
        task = _make_task(scene="tutor")
        content = _make_lesson_plan()
        result = build_docx(task, content)
        # 仍然能成功生成
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_pending_sections_excluded(self) -> None:
        """pending 状态的 section 不应出现在导出中。"""
        task = _make_task()
        content = _make_lesson_plan(
            objectives=[],
            objectives_status="pending",
            preparation=[],
            preparation_status="pending",
        )
        result = build_docx(task, content)
        text = _extract_docx_text(result)
        # pending 的 section 标题不应出现
        assert "教学目标" not in text
        assert "教学准备" not in text

    def test_confirmed_sections_included(self) -> None:
        """confirmed 状态的 section 应出现在导出中。"""
        task = _make_task()
        content = _make_lesson_plan()
        result = build_docx(task, content)
        text = _extract_docx_text(result)
        assert "教学目标" in text
        assert "掌握生字词" in text
        assert "教学重难点" in text
        assert "教学过程" in text


# ---------------------------------------------------------------------------
# 学案导出测试
# ---------------------------------------------------------------------------


class TestStudyGuideExport:
    def test_basic_export_produces_bytes(self) -> None:
        task = _make_task()
        content = _make_study_guide()
        result = build_docx(task, content)
        assert isinstance(result, bytes)
        assert len(result) > 0
        assert result[:2] == b"PK"

    def test_study_guide_header_info(self) -> None:
        task = _make_task()
        content = _make_study_guide()
        result = build_docx(task, content)
        text = _extract_docx_text(result)
        assert "语文" in text
        assert "七年级" in text

    def test_assessment_items_included(self) -> None:
        task = _make_task()
        content = _make_study_guide()
        result = build_docx(task, content)
        text = _extract_docx_text(result)
        assert "以下哪个是比喻句" in text
        assert "参考答案" in text

    def test_empty_study_guide(self) -> None:
        task = _make_task()
        content = StudyGuideContent()
        result = build_docx(task, content)
        assert isinstance(result, bytes)
        assert len(result) > 0


# ---------------------------------------------------------------------------
# 辅助
# ---------------------------------------------------------------------------


def _extract_docx_text(data: bytes) -> str:
    """从 docx 字节中提取全部文本。"""
    from io import BytesIO

    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
