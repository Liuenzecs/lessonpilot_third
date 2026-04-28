from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document as DocxDocument

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PPTX_MIME = "application/vnd.openxmlformats-officedocument.presentationml.presentation"


def _docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=1, cols=len(table_rows[0]))
        for col_index, value in enumerate(table_rows[0]):
            table.rows[0].cells[col_index].text = value
        for row_values in table_rows[1:]:
            row = table.add_row()
            for col_index, value in enumerate(row_values):
                row.cells[col_index].text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pptx_bytes() -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "ppt/slides/slide1.xml",
            """<p:sld xmlns:p="p" xmlns:a="a"><p:cSld><p:spTree>
            <a:t>《春》导入</a:t><a:t>展示春景图片</a:t>
            </p:spTree></p:cSld></p:sld>""",
        )
        archive.writestr(
            "ppt/slides/slide2.xml",
            """<p:sld xmlns:p="p" xmlns:a="a"><p:cSld><p:spTree>
            <a:t>品味语言</a:t><a:t>比喻拟人赏析</a:t>
            </p:spTree></p:cSld></p:sld>""",
        )
    return buffer.getvalue()


def _create_document(client, auth_headers, content: dict | None = None):
    task_response = client.post(
        "/api/v1/tasks/",
        headers=auth_headers,
        json={"subject": "语文", "grade": "七年级", "topic": "春", "lesson_type": "lesson_plan"},
    )
    task = task_response.json()
    document = client.get(f"/api/v1/documents/?task_id={task['id']}", headers=auth_headers).json()["items"][0]
    if content is not None:
        patched = client.patch(
            f"/api/v1/documents/{document['id']}",
            headers=auth_headers,
            json={"version": document["version"], "content": content},
        )
        document = patched.json()
    return task, document


def _ready_lesson_content() -> dict:
    return {
        "doc_type": "lesson_plan",
        "header": {
            "title": "春",
            "subject": "语文",
            "grade": "七年级",
            "classHour": 1,
            "lessonCategory": "new",
            "teacher": "张老师",
        },
        "objectives": [
            {"dimension": "knowledge", "content": "朗读课文，概括春景特点"},
            {"dimension": "ability", "content": "品味比喻拟人句的表达效果"},
        ],
        "objectives_status": "confirmed",
        "key_points": {"keyPoints": ["春景特点", "比喻拟人"], "difficulties": ["情景交融"]},
        "key_points_status": "confirmed",
        "preparation": ["朗读音频", "学习单"],
        "preparation_status": "confirmed",
        "teaching_process": [
            {
                "phase": "导入新课",
                "duration": 5,
                "teacher_activity": "展示春景图片，引出课题。",
                "student_activity": "观察图片，说出春景特点。",
                "design_intent": "激发阅读兴趣。",
                "status": "confirmed",
            },
            {
                "phase": "朗读感知",
                "duration": 15,
                "teacher_activity": "指导朗读，梳理春景特点。",
                "student_activity": "朗读课文，圈画关键词。",
                "design_intent": "落实整体感知。",
                "status": "confirmed",
            },
            {
                "phase": "品味语言",
                "duration": 20,
                "teacher_activity": "引导赏析比喻拟人句。",
                "student_activity": "小组讨论表达效果。",
                "design_intent": "突破语言赏析难点。",
                "status": "confirmed",
            },
        ],
        "teaching_process_status": "confirmed",
        "board_design": "春：盼春、绘春、赞春",
        "board_design_status": "confirmed",
        "reflection": "",
        "reflection_status": "pending",
        "section_references": {},
    }


def test_school_template_preview_confirm_export_and_isolation(client, auth_headers):
    template_bytes = _docx_bytes(
        [
            "七年级语文教案模板",
            "学科：语文 年级：七年级 课题：",
            "一、教学目标",
            "二、教学过程",
            "三、教学反思",
            "审核签字：",
        ],
        [["教学环节", "时间", "教师活动", "学生活动", "设计意图"], ["", "", "", "", ""]],
    )
    preview_response = client.post(
        "/api/v1/templates/school/preview",
        headers=auth_headers,
        files={"file": ("school-template.docx", template_bytes, DOCX_MIME)},
    )
    assert preview_response.status_code == 200
    preview = preview_response.json()
    assert "teaching_process" in preview["section_order"]
    assert preview["table_layouts"][0]["columns"] == ["教学环节", "时间", "教师活动", "学生活动", "设计意图"]

    confirm_response = client.post(
        "/api/v1/templates/school/confirm",
        headers=auth_headers,
        json={"preview": preview, "name": "本校语文教案模板"},
    )
    assert confirm_response.status_code == 201
    template_id = confirm_response.json()["id"]

    _task, document = _create_document(client, auth_headers, _ready_lesson_content())
    export_response = client.get(
        f"/api/v1/documents/{document['id']}/export?template_id={template_id}",
        headers=auth_headers,
    )
    assert export_response.status_code == 200
    text = _extract_docx_text(export_response.content)
    assert "教学环节" in text
    assert "教师活动" in text
    assert "审核签字" in text

    other = client.post(
        "/api/v1/auth/register",
        json={"email": "phase7-other@example.com", "name": "Other", "password": "Password123"},
    )
    other_headers = {"Authorization": f"Bearer {other.json()['access_token']}"}
    assert client.get("/api/v1/templates/school/personal", headers=other_headers).json() == []
    forbidden_export = client.get(
        f"/api/v1/documents/{document['id']}/export?template_id={template_id}",
        headers=other_headers,
    )
    assert forbidden_export.status_code == 404


def test_quality_check_v2_reports_alignment_and_actionable_warnings(client, auth_headers):
    bad_content = _ready_lesson_content()
    bad_content["objectives"] = [{"dimension": "knowledge", "content": "提高综合素养"}]
    bad_content["key_points"] = {"keyPoints": ["文言虚词"], "difficulties": ["倒装句式"]}
    for step in bad_content["teaching_process"]:
        step["student_activity"] = "认真听讲"
        step["teacher_activity"] = "讲解课文"
        step["design_intent"] = "完成教学任务"
    _task, document = _create_document(client, auth_headers, bad_content)

    response = client.post(f"/api/v1/documents/{document['id']}/quality-check", headers=auth_headers)

    payload = response.json()
    assert payload["readiness"] == "needs_fixes"
    assert payload["alignment_map"]
    warning_text = response.text
    assert "表述偏空泛" in warning_text
    assert "教学过程没有明显承接教学目标" in warning_text
    assert "学生活动整体偏被动" in warning_text


def test_teaching_package_generation_requires_confirmed_lesson_and_creates_pending_package(client, auth_headers):
    _task, empty_document = _create_document(client, auth_headers)
    blocked = client.post(f"/api/v1/documents/{empty_document['id']}/teaching-package", headers=auth_headers)
    assert blocked.status_code == 400
    assert "请先确认" in blocked.text

    _task, document = _create_document(client, auth_headers, _ready_lesson_content())
    response = client.post(f"/api/v1/documents/{document['id']}/teaching-package", headers=auth_headers)

    assert response.status_code == 201
    payload = response.json()
    assert payload["status"] == "pending"
    assert payload["content"]["study_guide"]["learning_objectives_status"] == "pending"
    assert len(payload["content"]["ppt_outline"]) >= 3
    assert payload["content"]["talk_script"]["questions"]


def test_personal_asset_docx_and_pptx_preview_confirm_are_private(client, auth_headers):
    docx_preview = client.post(
        "/api/v1/personal-assets/preview",
        headers=auth_headers,
        files={
            "file": (
                "old-lesson.docx",
                _docx_bytes(["《春》教案", "一、教学目标", "朗读课文", "二、教学过程", "导入、朗读、赏析"]),
                DOCX_MIME,
            )
        },
    )
    assert docx_preview.status_code == 200
    assert docx_preview.json()["asset_type"] == "lesson_plan"

    asset_response = client.post(
        "/api/v1/personal-assets/confirm",
        headers=auth_headers,
        json={"preview": docx_preview.json()},
    )
    assert asset_response.status_code == 201
    assert asset_response.json()["source_filename"] == "old-lesson.docx"

    pptx_preview = client.post(
        "/api/v1/personal-assets/preview",
        headers=auth_headers,
        files={"file": ("spring.pptx", _pptx_bytes(), PPTX_MIME)},
    )
    assert pptx_preview.status_code == 200
    assert pptx_preview.json()["asset_type"] == "ppt_outline"
    assert "品味语言" in pptx_preview.text

    other = client.post(
        "/api/v1/auth/register",
        json={"email": "phase10-other@example.com", "name": "Other", "password": "Password123"},
    )
    other_headers = {"Authorization": f"Bearer {other.json()['access_token']}"}
    assert client.get("/api/v1/personal-assets/", headers=other_headers).json() == []


def _extract_docx_text(data: bytes) -> str:
    document = DocxDocument(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
