from __future__ import annotations

import json
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document as DocxDocument

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PPTX_MIME = "application/vnd.openxmlformats-officedocument.presentationml.presentation"


def _parse_sse_events(raw_text: str) -> list[tuple[str, dict]]:
    events: list[tuple[str, dict]] = []
    for chunk in raw_text.strip().split("\n\n"):
        if not chunk.strip():
            continue
        event_name = ""
        data_payload = ""
        for line in chunk.splitlines():
            if line.startswith("event: "):
                event_name = line.removeprefix("event: ").strip()
            if line.startswith("data: "):
                data_payload = line.removeprefix("data: ").strip()
        if event_name and data_payload:
            events.append((event_name, json.loads(data_payload)))
    return events


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("《春》旧教案")
    document.add_paragraph("学科：语文 年级：七年级 课题：春")
    document.add_paragraph("一、教学目标")
    document.add_paragraph("朗读课文，圈画春草图、春花图中的关键词。")
    document.add_paragraph("二、教学过程")
    document.add_paragraph("品味比喻拟人，组织学生小组展示。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pptx_bytes() -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "ppt/slides/slide1.xml",
            """<p:sld xmlns:p="p" xmlns:a="a"><p:cSld><p:spTree>
            <a:t>《春》朗读导入</a:t><a:t>春草图、春花图任务</a:t>
            </p:spTree></p:cSld></p:sld>""",
        )
    return buffer.getvalue()


def _create_asset(client, auth_headers, *, file_name: str = "spring.docx") -> dict:
    file_bytes = _docx_bytes() if file_name.endswith(".docx") else _pptx_bytes()
    mime_type = DOCX_MIME if file_name.endswith(".docx") else PPTX_MIME
    preview = client.post(
        "/api/v1/personal-assets/preview",
        headers=auth_headers,
        files={"file": (file_name, file_bytes, mime_type)},
    )
    assert preview.status_code == 200
    response = client.post(
        "/api/v1/personal-assets/confirm",
        headers=auth_headers,
        json={"preview": preview.json(), "title": "春的旧资料"},
    )
    assert response.status_code == 201
    return response.json()


def _create_task(client, auth_headers, topic: str = "春") -> dict:
    response = client.post(
        "/api/v1/tasks/",
        headers=auth_headers,
        json={"subject": "语文", "grade": "七年级", "topic": topic, "lesson_type": "lesson_plan"},
    )
    assert response.status_code == 201
    return response.json()


def _ready_lesson_content() -> dict:
    return {
        "doc_type": "lesson_plan",
        "header": {
            "title": "春",
            "subject": "语文",
            "grade": "七年级",
            "classHour": 1,
            "lessonCategory": "new",
            "teacher": "",
        },
        "objectives": [{"dimension": "knowledge", "content": "提高综合素养"}],
        "objectives_status": "confirmed",
        "key_points": {"keyPoints": ["比喻拟人"], "difficulties": ["情景交融"]},
        "key_points_status": "confirmed",
        "preparation": ["课件"],
        "preparation_status": "confirmed",
        "teaching_process": [
            {
                "phase": "导入",
                "duration": 5,
                "teacher_activity": "讲解课文",
                "student_activity": "认真听讲",
                "design_intent": "完成教学任务",
                "status": "confirmed",
            },
            {
                "phase": "精读",
                "duration": 25,
                "teacher_activity": "讲解课文",
                "student_activity": "认真听讲",
                "design_intent": "完成教学任务",
                "status": "confirmed",
            },
            {
                "phase": "小结",
                "duration": 10,
                "teacher_activity": "讲解课文",
                "student_activity": "认真听讲",
                "design_intent": "完成教学任务",
                "status": "confirmed",
            },
        ],
        "teaching_process_status": "confirmed",
        "board_design": "春：绘春",
        "board_design_status": "confirmed",
        "reflection": "",
        "reflection_status": "pending",
        "section_references": {},
    }


def test_personal_asset_recommendation_is_private(client, auth_headers):
    asset = _create_asset(client, auth_headers)

    response = client.get(
        "/api/v1/personal-assets/recommend?subject=语文&grade=七年级&topic=春",
        headers=auth_headers,
    )
    assert response.status_code == 200
    assert response.json()[0]["asset_id"] == asset["id"]
    assert "春" in response.text
    unmatched = client.get(
        "/api/v1/personal-assets/recommend?subject=语文&grade=七年级&topic=桃花源记",
        headers=auth_headers,
    )
    assert unmatched.json() == []

    other = client.post(
        "/api/v1/auth/register",
        json={"email": "phase11-other@example.com", "name": "Other", "password": "Password123"},
    )
    other_headers = {"Authorization": f"Bearer {other.json()['access_token']}"}
    other_response = client.get(
        "/api/v1/personal-assets/recommend?subject=语文&grade=七年级&topic=春",
        headers=other_headers,
    )
    assert other_response.json() == []


def test_generation_uses_personal_asset_status_and_section_references(client, auth_headers):
    asset = _create_asset(client, auth_headers, file_name="spring.pptx")
    task = _create_task(client, auth_headers)
    start_response = client.post(
        f"/api/v1/tasks/{task['id']}/generate",
        headers=auth_headers,
        json={"use_personal_assets": True, "personal_asset_ids": [asset["id"]]},
    )
    assert start_response.status_code == 200

    with client.stream("GET", start_response.json()["stream_url"], headers=auth_headers) as resp:
        events = _parse_sse_events("".join(resp.iter_text()))

    asset_events = [payload for event, payload in events if event == "asset_status"]
    assert asset_events[0]["status"] == "ready"
    assert asset_events[0]["matched_assets"][0]["asset_id"] == asset["id"]

    documents = client.get(f"/api/v1/documents/?task_id={task['id']}", headers=auth_headers).json()["items"]
    references = documents[0]["content"]["section_references"]
    assert references["objectives"][0]["knowledge_type"] == "personal_asset"
    assert references["teaching_process"][0]["source"].startswith("我的资料库")


def test_generation_rejects_other_users_personal_asset(client, auth_headers):
    asset = _create_asset(client, auth_headers)
    other = client.post(
        "/api/v1/auth/register",
        json={"email": "phase11-generate-other@example.com", "name": "Other", "password": "Password123"},
    )
    other_headers = {"Authorization": f"Bearer {other.json()['access_token']}"}
    task = _create_task(client, other_headers)

    response = client.post(
        f"/api/v1/tasks/{task['id']}/generate",
        headers=other_headers,
        json={"use_personal_assets": True, "personal_asset_ids": [asset["id"]]},
    )

    assert response.status_code == 404


def test_quality_fix_creates_pending_targeted_revisions(client, auth_headers):
    task = _create_task(client, auth_headers)
    document = client.get(f"/api/v1/documents/?task_id={task['id']}", headers=auth_headers).json()["items"][0]
    patched = client.patch(
        f"/api/v1/documents/{document['id']}",
        headers=auth_headers,
        json={"version": document["version"], "content": _ready_lesson_content()},
    ).json()

    objective_fix = client.post(
        f"/api/v1/documents/{patched['id']}/quality-fix",
        headers=auth_headers,
        json={"section": "objectives", "message": "教学目标第 1 条表述偏空泛。", "suggestion": "改具体。"},
    )
    assert objective_fix.status_code == 200
    content = objective_fix.json()["content"]
    assert content["objectives_status"] == "pending"
    assert "提高综合素养" not in content["objectives"][0]["content"]

    process_fix = client.post(
        f"/api/v1/documents/{patched['id']}/quality-fix",
        headers=auth_headers,
        json={"section": "teaching_process", "message": "学生活动整体偏被动。", "suggestion": "补充任务。"},
    )
    assert process_fix.status_code == 200
    process_content = process_fix.json()["content"]
    assert process_content["teaching_process_status"] == "pending"
    assert "朗读课文" in process_content["teaching_process"][0]["student_activity"]

    key_point_fix = client.post(
        f"/api/v1/documents/{patched['id']}/quality-fix",
        headers=auth_headers,
        json={"section": "key_points", "message": "教学重难点没有在教学过程中明显展开。", "suggestion": "写进过程。"},
    )
    assert key_point_fix.status_code == 200
    key_point_content = key_point_fix.json()["content"]
    assert key_point_content["teaching_process_status"] == "pending"
    assert "比喻拟人" in key_point_content["teaching_process"][0]["teacher_activity"]
