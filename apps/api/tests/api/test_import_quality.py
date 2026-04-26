from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _lesson_docx_bytes(*, include_unmapped: bool = True) -> bytes:
    document = DocxDocument()
    document.add_paragraph("《春》教案")
    document.add_paragraph("学科：语文 年级：七年级 课题：春 课时：1")
    if include_unmapped:
        document.add_paragraph("课堂亮点：朗读节奏要轻快。")
    document.add_paragraph("一、教学目标")
    document.add_paragraph("1. 朗读课文，感受春天景物特点。")
    document.add_paragraph("2. 学习比喻、拟人等修辞方法。")
    document.add_paragraph("二、教学重难点")
    document.add_paragraph("教学重点：品味课文中描写春景的语言。")
    document.add_paragraph("教学难点：理解情景交融的表达效果。")
    document.add_paragraph("三、教学准备")
    document.add_paragraph("多媒体课件")
    document.add_paragraph("朗读音频")
    document.add_paragraph("四、教学过程")
    table = document.add_table(rows=1, cols=5)
    headers = ["教学环节", "时长", "教师活动", "学生活动", "设计意图"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    rows = [
        ["导入新课", "5分钟", "展示春景图片，引出课题。", "观察图片，说出春天印象。", "激发阅读兴趣。"],
        ["初读感知", "15分钟", "范读并指导停顿。", "圈画关键词，交流初读感受。", "整体感知文章。"],
        ["品味语言", "20分钟", "引导赏析比喻拟人句。", "小组讨论并汇报。", "落实语言训练。"],
    ]
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].text = value
    document.add_paragraph("五、板书设计")
    document.add_paragraph("春：盼春、绘春、赞春")
    document.add_paragraph("六、教学反思")
    document.add_paragraph("课后填写。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_lesson_plan(client, auth_headers, content_override: dict | None = None):
    create_response = client.post(
        "/api/v1/tasks/",
        headers=auth_headers,
        json={"subject": "语文", "grade": "七年级", "topic": "春", "lesson_type": "lesson_plan"},
    )
    task = create_response.json()
    document = client.get(f"/api/v1/documents/?task_id={task['id']}", headers=auth_headers).json()["items"][0]
    if content_override is not None:
        patch_response = client.patch(
            f"/api/v1/documents/{document['id']}",
            headers=auth_headers,
            json={"version": document["version"], "content": content_override},
        )
        document = patch_response.json()
    return task, document


def test_lesson_plan_import_preview_maps_core_sections(client, auth_headers):
    response = client.post(
        "/api/v1/import/lesson-plan/preview",
        headers=auth_headers,
        files={
            "file": (
                "spring.docx",
                _lesson_docx_bytes(),
                DOCX_MIME,
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["metadata"]["topic"] == "春"
    assert set(payload["mapped_sections"]) >= {
        "objectives",
        "key_points",
        "preparation",
        "teaching_process",
        "board_design",
        "reflection",
    }
    content = payload["content"]
    assert content["objectives_status"] == "pending"
    assert content["teaching_process_status"] == "pending"
    assert len(content["teaching_process"]) == 3
    assert content["teaching_process"][0]["student_activity"] == "观察图片，说出春天印象。"
    assert payload["unmapped_sections"][0]["content"] == "课堂亮点：朗读节奏要轻快。"


def test_lesson_plan_import_confirm_creates_pending_document(client, auth_headers):
    preview = client.post(
        "/api/v1/import/lesson-plan/preview",
        headers=auth_headers,
        files={"file": ("spring.docx", _lesson_docx_bytes(), DOCX_MIME)},
    ).json()

    response = client.post(
        "/api/v1/import/lesson-plan/confirm",
        headers=auth_headers,
        json={"metadata": preview["metadata"], "content": preview["content"]},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["task"]["status"] == "imported"
    assert payload["task"]["lesson_type"] == "lesson_plan"
    assert payload["document"]["content"]["objectives_status"] == "pending"
    assert payload["document"]["content"]["teaching_process_status"] == "pending"


def test_lesson_plan_import_rejects_unsupported_or_large_files(client, auth_headers):
    unsupported = client.post(
        "/api/v1/import/lesson-plan/preview",
        headers=auth_headers,
        files={"file": ("old.doc", b"not-docx", "application/msword")},
    )
    assert unsupported.status_code == 400

    too_large = client.post(
        "/api/v1/import/lesson-plan/preview",
        headers=auth_headers,
        files={"file": ("large.docx", b"x" * (5 * 1024 * 1024 + 1), DOCX_MIME)},
    )
    assert too_large.status_code == 413


def test_quality_check_reports_blocked_needs_fixes_and_ready(client, auth_headers):
    _task, empty_document = _create_lesson_plan(client, auth_headers)
    blocked = client.post(
        f"/api/v1/documents/{empty_document['id']}/quality-check",
        headers=auth_headers,
    )
    assert blocked.status_code == 200
    assert blocked.json()["readiness"] == "blocked"

    preview = client.post(
        "/api/v1/import/lesson-plan/preview",
        headers=auth_headers,
        files={"file": ("spring.docx", _lesson_docx_bytes(include_unmapped=False), DOCX_MIME)},
    ).json()
    imported = client.post(
        "/api/v1/import/lesson-plan/confirm",
        headers=auth_headers,
        json={"metadata": preview["metadata"], "content": preview["content"]},
    ).json()["document"]
    needs_fixes = client.post(
        f"/api/v1/documents/{imported['id']}/quality-check",
        headers=auth_headers,
    )
    assert needs_fixes.json()["readiness"] == "needs_fixes"
    assert "待确认" in needs_fixes.text

    ready_content = imported["content"]
    ready_content["objectives_status"] = "confirmed"
    ready_content["key_points_status"] = "confirmed"
    ready_content["preparation_status"] = "confirmed"
    ready_content["teaching_process_status"] = "confirmed"
    ready_content["board_design_status"] = "confirmed"
    ready_content["reflection_status"] = "confirmed"
    _task, ready_document = _create_lesson_plan(client, auth_headers, ready_content)
    ready = client.post(
        f"/api/v1/documents/{ready_document['id']}/quality-check",
        headers=auth_headers,
    )
    assert ready.json()["readiness"] == "ready"


def test_quality_check_is_user_isolated(client, auth_headers):
    _task, document = _create_lesson_plan(client, auth_headers)
    register = client.post(
        "/api/v1/auth/register",
        json={"email": "other@example.com", "name": "Other", "password": "Password123"},
    )
    other_headers = {"Authorization": f"Bearer {register.json()['access_token']}"}

    response = client.post(f"/api/v1/documents/{document['id']}/quality-check", headers=other_headers)

    assert response.status_code == 404
