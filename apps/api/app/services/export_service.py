"""教案 / 学案 Word 导出服务。

基于 LessonPlanContent / StudyGuideContent 结构化模型，
生成学校标准格式 Word 文档（.docx）。
"""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models import Task
from app.schemas.content import (
    AssessmentItem,
    DocumentContent,
    LessonPlanContent,
    StudyGuideContent,
    TeachingProcessStep,
    is_lesson_plan,
    is_study_guide,
)
from app.services.word_formula import build_omath, split_formula_segments

# ---------------------------------------------------------------------------
# 颜色常量（中式现代风）
# ---------------------------------------------------------------------------

_COLOR_INK = "2c2c2c"       # 墨色 — 正文
_COLOR_STONE = "3a7ca5"     # 石青 — 标题/强调
_COLOR_MUTED = "8c8c8c"     # 灰色 — 次要信息
_COLOR_BORDER = "d0ccc4"    # 象牙边框

# ---------------------------------------------------------------------------
# 中文字号 → Pt 映射（近似）
# ---------------------------------------------------------------------------

_PT_ZERO = Pt(0)
_PT_TWO = Pt(2)
_PT_FOUR = Pt(4)
_PT_SIX = Pt(6)
_PT_EIGHT = Pt(8)
_PT_TWELVE = Pt(12)
_PT_FOURTEEN = Pt(14)
_PT_SIXTEEN = Pt(16)
_PT_EIGHTEEN = Pt(18)
_PT_TWENTYFOUR = Pt(24)


# ---------------------------------------------------------------------------
# 样式配置
# ---------------------------------------------------------------------------

def _set_font(
    style_or_run,
    name: str,
    size: Pt,
    *,
    bold: bool = False,
    color: str | None = None,
) -> None:
    font = style_or_run.font if hasattr(style_or_run, "font") else style_or_run
    font.name = name
    font.size = size
    font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)

    # 设置东亚字体
    r_pr = style_or_run._element.get_or_add_rPr() if hasattr(style_or_run, "_element") else None
    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), name)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)


def _configure_page(document: DocxDocument) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = document.styles["Normal"]
    _set_font(normal, "宋体", _PT_TWELVE, color=_COLOR_INK)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = _PT_FOUR


# ---------------------------------------------------------------------------
# 段落辅助
# ---------------------------------------------------------------------------

def _add_inline_content(
    paragraph,
    text: str,
    *,
    font_name: str = "宋体",
    font_size: Pt = _PT_TWELVE,
    bold: bool = False,
    color: str | None = None,
) -> None:
    for segment in split_formula_segments(text):
        if segment.kind == "formula":
            paragraph._element.append(build_omath(segment.text))
        elif segment.text:
            run = paragraph.add_run(segment.text)
            _set_font(run, font_name, font_size, bold=bold, color=color)


def _add_paragraph(
    document: DocxDocument,
    text: str,
    *,
    font_name: str = "宋体",
    font_size: Pt = _PT_TWELVE,
    bold: bool = False,
    color: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_before: Pt = _PT_ZERO,
    space_after: Pt = _PT_FOUR,
    keep_with_next: bool = False,
) -> None:
    p = document.add_paragraph()
    _add_inline_content(
        p,
        text,
        font_name=font_name,
        font_size=font_size,
        bold=bold,
        color=color,
    )
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if keep_with_next:
        p.paragraph_format.keep_with_next = True


def _add_section_title(document: DocxDocument, title: str) -> None:
    _add_paragraph(
        document,
        title,
        font_name="微软雅黑",
        font_size=_PT_FOURTEEN,
        bold=True,
        color=_COLOR_STONE,
        space_before=_PT_SIXTEEN,
        space_after=_PT_EIGHT,
        keep_with_next=True,
    )


def _add_body(document: DocxDocument, text: str, *, indent: Pt | None = None) -> None:
    p = document.add_paragraph()
    _add_inline_content(p, text, font_name="宋体", font_size=_PT_TWELVE, color=_COLOR_INK)
    if indent is not None:
        p.paragraph_format.left_indent = indent


def _add_list_items(document: DocxDocument, items: list[str], *, prefix: str = "• ") -> None:
    for item in items:
        _add_body(document, f"{prefix}{item}", indent=_PT_EIGHTEEN)


def _add_empty_area(document: DocxDocument, hint: str = "（请在此处填写）") -> None:
    p = document.add_paragraph()
    run = p.add_run(hint)
    _set_font(run, "宋体", _PT_TWELVE, color=_COLOR_MUTED)
    p.paragraph_format.space_before = _PT_EIGHT
    p.paragraph_format.space_after = _PT_TWENTYFOUR
    # 添加下划线区域
    for _ in range(3):
        line_p = document.add_paragraph()
        run = line_p.add_run("　")
        _set_font(run, "宋体", _PT_TWELVE)
        line_p.paragraph_format.space_after = _PT_TWO
        pPr = line_p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), _COLOR_BORDER)
        pBdr.append(bottom)
        pPr.append(pBdr)


# ---------------------------------------------------------------------------
# 表格辅助
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str, *, bold: bool = False, font_size: Pt = _PT_TWELVE) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    _add_inline_content(p, text, font_name="宋体", font_size=font_size, bold=bold, color=_COLOR_INK)
    p.paragraph_format.space_before = _PT_TWO
    p.paragraph_format.space_after = _PT_TWO
    p.paragraph_format.line_spacing = 1.3


def _set_cell_shading(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), _COLOR_BORDER)
        borders.append(border)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# 教案导出
# ---------------------------------------------------------------------------

_OBJECTIVE_DIMENSION_LABELS = {
    "knowledge": "知识与技能",
    "ability": "过程与方法",
    "emotion": "情感态度与价值观",
}

_LESSON_CATEGORY_LABELS = {
    "new": "新授课",
    "review": "复习课",
    "exercise": "练习课",
    "comprehensive": "综合课",
}


def _render_lesson_plan_header(
    document: DocxDocument,
    content: LessonPlanContent,
    task: Task,
) -> None:
    header = content.header
    # 大标题
    _add_paragraph(
        document,
        header.title or task.title,
        font_name="微软雅黑",
        font_size=_PT_EIGHTEEN,
        bold=True,
        color=_COLOR_INK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=_PT_SIX,
    )
    # 元信息
    category_label = _LESSON_CATEGORY_LABELS.get(task.lesson_category, task.lesson_category)
    meta_parts = [task.subject, task.grade, f"第{task.class_hour}课时", category_label]
    if header.teacher:
        meta_parts.append(f"授课教师：{header.teacher}")
    _add_paragraph(
        document,
        "　｜　".join(meta_parts),
        font_size=_PT_TWELVE,
        color=_COLOR_MUTED,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=_PT_TWELVE,
    )


def _render_objectives(document: DocxDocument, content: LessonPlanContent) -> None:
    if content.objectives_status != "confirmed" or not content.objectives:
        return
    _add_section_title(document, "一、教学目标")
    for obj in content.objectives:
        dimension = _OBJECTIVE_DIMENSION_LABELS.get(obj.dimension, obj.dimension)
        _add_body(document, f"【{dimension}】{obj.content}", indent=_PT_EIGHTEEN)


def _render_key_points(document: DocxDocument, content: LessonPlanContent) -> None:
    if content.key_points_status != "confirmed":
        return
    kp = content.key_points
    if not kp.key_points and not kp.difficulties:
        return
    _add_section_title(document, "二、教学重难点")
    if kp.key_points:
        _add_paragraph(
            document,
            "教学重点：",
            font_name="微软雅黑",
            font_size=_PT_TWELVE,
            bold=True,
            color=_COLOR_INK,
            space_before=_PT_FOUR,
            space_after=_PT_TWO,
            keep_with_next=True,
        )
        _add_list_items(document, kp.key_points)
    if kp.difficulties:
        _add_paragraph(
            document,
            "教学难点：",
            font_name="微软雅黑",
            font_size=_PT_TWELVE,
            bold=True,
            color=_COLOR_INK,
            space_before=_PT_FOUR,
            space_after=_PT_TWO,
            keep_with_next=True,
        )
        _add_list_items(document, kp.difficulties)


def _render_preparation(document: DocxDocument, content: LessonPlanContent) -> None:
    if content.preparation_status != "confirmed" or not content.preparation:
        return
    _add_section_title(document, "三、教学准备")
    _add_list_items(document, content.preparation)


def _render_teaching_process(
    document: DocxDocument,
    content: LessonPlanContent,
    scene: str,
) -> None:
    if content.teaching_process_status != "confirmed" or not content.teaching_process:
        return
    _add_section_title(document, "四、教学过程")

    # 公立校和机构使用完整 5 列表格，家教简化为 4 列（省略设计意图）
    if scene == "tutor":
        headers = ["教学环节", "时长", "教师活动", "学生活动"]
        col_widths = [Cm(2.5), Cm(1.5), Cm(5.5), Cm(5.5)]
    else:
        headers = ["教学环节", "时长", "教师活动", "学生活动", "设计意图"]
        col_widths = [Cm(2.2), Cm(1.3), Cm(5.0), Cm(4.5), Cm(3.0)]

    num_cols = len(headers)
    table = document.add_table(rows=1 + len(content.teaching_process), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # 表头行
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, header_text, bold=True)
        _set_cell_shading(cell, "f0ebe0")
    # 设置列宽
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # 数据行
    for row_idx, step in enumerate(content.teaching_process):
        row = table.rows[row_idx + 1]
        _set_cell_text(row.cells[0], step.phase or f"环节{row_idx + 1}")
        _set_cell_text(row.cells[1], f"{step.duration}分钟")
        _set_cell_text(row.cells[2], step.teacher_activity)
        _set_cell_text(row.cells[3], step.student_activity)
        if num_cols >= 5:
            _set_cell_text(row.cells[4], step.design_intent)

    # 表格后空行
    document.add_paragraph()


def _render_board_design(document: DocxDocument, content: LessonPlanContent) -> None:
    if content.board_design_status != "confirmed" or not content.board_design:
        return
    _add_section_title(document, "五、板书设计")
    _add_body(document, content.board_design)


def _render_reflection(document: DocxDocument, content: LessonPlanContent) -> None:
    if content.reflection_status != "confirmed" and not content.reflection:
        # 反思区即使没有 confirmed 也保留留空区域
        pass
    _add_section_title(document, "六、教学反思")
    if content.reflection:
        _add_body(document, content.reflection)
    else:
        _add_empty_area(document, "（课后填写教学反思）")


def _build_lesson_plan_docx(
    document: DocxDocument,
    task: Task,
    content: LessonPlanContent,
) -> None:
    scene = task.scene
    _render_lesson_plan_header(document, content, task)
    _render_objectives(document, content)
    _render_key_points(document, content)
    _render_preparation(document, content)
    _render_teaching_process(document, content, scene)
    _render_board_design(document, content)
    _render_reflection(document, content)


def _build_templated_lesson_plan_docx(
    document: DocxDocument,
    task: Task,
    content: LessonPlanContent,
    template_spec: dict,
) -> None:
    _render_template_header(document, content, task, template_spec)
    section_order = template_spec.get("section_order") or [
        "objectives",
        "key_points",
        "preparation",
        "teaching_process",
        "board_design",
        "reflection",
    ]
    rendered = set()
    for section_name in section_order:
        _render_lesson_section_by_name(document, task, content, section_name, template_spec)
        rendered.add(section_name)
    for section_name in ["objectives", "key_points", "preparation", "teaching_process", "board_design", "reflection"]:
        if section_name not in rendered:
            _render_lesson_section_by_name(document, task, content, section_name, template_spec)
    for blank_label in template_spec.get("blank_areas") or []:
        if blank_label not in {"教学反思", "课后反思", "教后反思"}:
            _add_section_title(document, str(blank_label))
            _add_empty_area(document, f"（{blank_label}）")


def _render_template_header(
    document: DocxDocument,
    content: LessonPlanContent,
    task: Task,
    template_spec: dict,
) -> None:
    _add_paragraph(
        document,
        content.header.title or task.title,
        font_name="微软雅黑",
        font_size=_PT_EIGHTEEN,
        bold=True,
        color=_COLOR_INK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=_PT_SIX,
    )
    metadata_labels = _template_metadata_labels(template_spec)
    if metadata_labels:
        table = document.add_table(rows=0, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table)
        for left, right in _pairwise(metadata_labels):
            row = table.add_row()
            _set_cell_text(row.cells[0], left, bold=True)
            _set_cell_shading(row.cells[0], "f0ebe0")
            _set_cell_text(row.cells[1], _metadata_value(left, content, task))
            _set_cell_text(row.cells[2], right, bold=True)
            _set_cell_shading(row.cells[2], "f0ebe0")
            _set_cell_text(row.cells[3], _metadata_value(right, content, task))
        document.add_paragraph()
    else:
        _render_lesson_plan_header(document, content, task)


def _render_lesson_section_by_name(
    document: DocxDocument,
    task: Task,
    content: LessonPlanContent,
    section_name: str,
    template_spec: dict,
) -> None:
    if section_name == "objectives":
        _render_objectives(document, content)
    elif section_name == "key_points":
        _render_key_points(document, content)
    elif section_name == "preparation":
        _render_preparation(document, content)
    elif section_name == "teaching_process":
        _render_template_teaching_process(document, content, task.scene, template_spec)
    elif section_name == "board_design":
        _render_board_design(document, content)
    elif section_name == "reflection":
        _render_reflection(document, content)


def _render_template_teaching_process(
    document: DocxDocument,
    content: LessonPlanContent,
    scene: str,
    template_spec: dict,
) -> None:
    if content.teaching_process_status != "confirmed" or not content.teaching_process:
        return
    _add_section_title(document, "教学过程")
    headers = _template_process_columns(template_spec)
    if not headers:
        headers = (
            ["教学环节", "时长", "教师活动", "学生活动"]
            if scene == "tutor"
            else ["教学环节", "时长", "教师活动", "学生活动", "设计意图"]
        )
    table = document.add_table(rows=1 + len(content.teaching_process), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)
        _set_cell_shading(table.rows[0].cells[index], "f0ebe0")
    for row_index, step in enumerate(content.teaching_process, start=1):
        row = table.rows[row_index]
        for col_index, header in enumerate(headers):
            _set_cell_text(row.cells[col_index], _process_value(header, step))
    document.add_paragraph()


def _template_metadata_labels(template_spec: dict) -> list[str]:
    for layout in template_spec.get("table_layouts") or []:
        if layout.get("name") == "metadata":
            return [str(item) for item in layout.get("columns") or []][:8]
    labels = []
    for mapping in template_spec.get("field_mappings") or []:
        if str(mapping.get("content_field", "")).startswith("header"):
            labels.append(str(mapping.get("template_label") or ""))
    return [label for label in labels if label][:8]


def _template_process_columns(template_spec: dict) -> list[str]:
    for layout in template_spec.get("table_layouts") or []:
        if layout.get("name") == "teaching_process":
            return [str(item) for item in layout.get("columns") or [] if str(item).strip()]
    return []


def _metadata_value(label: str, content: LessonPlanContent, task: Task) -> str:
    if "课题" in label or "课名" in label or "教学内容" in label:
        return content.header.title or task.topic
    if "学科" in label:
        return content.header.subject or task.subject
    if "年级" in label:
        return content.header.grade or task.grade
    if "课时" in label:
        return f"第{task.class_hour}课时"
    if "课型" in label:
        return _LESSON_CATEGORY_LABELS.get(task.lesson_category, task.lesson_category)
    if "教师" in label:
        return content.header.teacher
    return ""


def _process_value(label: str, step: TeachingProcessStep) -> str:
    if "环节" in label or "流程" in label or "步骤" in label:
        return step.phase
    if "时" in label or "分钟" in label:
        return f"{step.duration}分钟"
    if "教师" in label or "教法" in label:
        return step.teacher_activity
    if "学生" in label or "学法" in label:
        return step.student_activity
    if "意图" in label or "说明" in label:
        return step.design_intent
    return ""


def _pairwise(labels: list[str]) -> list[tuple[str, str]]:
    result: list[tuple[str, str]] = []
    for index in range(0, len(labels), 2):
        left = labels[index]
        right = labels[index + 1] if index + 1 < len(labels) else ""
        result.append((left, right))
    return result


# ---------------------------------------------------------------------------
# 学案导出
# ---------------------------------------------------------------------------

def _render_study_guide_header(
    document: DocxDocument,
    content: StudyGuideContent,
    task: Task,
) -> None:
    header = content.header
    _add_paragraph(
        document,
        header.title or task.title,
        font_name="微软雅黑",
        font_size=_PT_EIGHTEEN,
        bold=True,
        color=_COLOR_INK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=_PT_EIGHT,
    )

    # 学生信息表格（2 列 × 3 行）
    table = document.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    info_pairs = [
        ("学科", header.subject or task.subject, "年级", header.grade or task.grade),
        ("班级", header.class_name or "", "姓名", header.student_name or ""),
        ("日期", header.date or "", "课题", header.title or task.topic),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(info_pairs):
        row = table.rows[row_idx]
        _set_cell_text(row.cells[0], l1, bold=True)
        _set_cell_shading(row.cells[0], "f0ebe0")
        _set_cell_text(row.cells[1], v1)
        _set_cell_text(row.cells[2], l2, bold=True)
        _set_cell_shading(row.cells[2], "f0ebe0")
        _set_cell_text(row.cells[3], v2)

    document.add_paragraph()


def _render_learning_objectives(document: DocxDocument, content: StudyGuideContent) -> None:
    if content.learning_objectives_status != "confirmed" or not content.learning_objectives:
        return
    _add_section_title(document, "一、学习目标")
    for obj in content.learning_objectives:
        _add_body(document, f"• {obj}", indent=_PT_EIGHTEEN)


def _render_key_difficulties(document: DocxDocument, content: StudyGuideContent) -> None:
    if content.key_difficulties_status != "confirmed" or not content.key_difficulties:
        return
    _add_section_title(document, "二、重点难点预测")
    _add_list_items(document, content.key_difficulties)


def _render_prior_knowledge(document: DocxDocument, content: StudyGuideContent) -> None:
    if content.prior_knowledge_status != "confirmed" or not content.prior_knowledge:
        return
    _add_section_title(document, "三、知识链接")
    _add_list_items(document, content.prior_knowledge)


def _render_assessment_items(
    document: DocxDocument,
    title: str,
    items: list[AssessmentItem],
) -> None:
    if not items:
        return
    _add_section_title(document, title)
    for idx, item in enumerate(items, start=1):
        type_labels = {"choice": "选择题", "fill_blank": "填空题", "short_answer": "简答题"}
        type_label = type_labels.get(item.item_type, "题")
        level_label = f"（{item.level}级）" if item.level else ""
        _add_paragraph(
            document,
            f"{idx}. [{type_label}{level_label}] {item.prompt}",
            font_size=_PT_TWELVE,
            space_before=_PT_FOUR,
            space_after=_PT_TWO,
            keep_with_next=True,
        )
        # 选项（选择题）
        for opt_idx, option in enumerate(item.options):
            prefix = chr(65 + opt_idx) if opt_idx < 26 else str(opt_idx + 1)
            _add_body(document, f"  {prefix}. {option}", indent=Pt(36))  # noqa: B008
        # 答案和解析
        if item.answer:
            _add_body(document, f"参考答案：{item.answer}", indent=Pt(36))  # noqa: B008
        if item.analysis:
            _add_body(document, f"解析：{item.analysis}", indent=Pt(36))  # noqa: B008


def _render_learning_process(document: DocxDocument, content: StudyGuideContent) -> None:
    lp = content.learning_process
    has_any = (
        (content.self_study_status == "confirmed" and lp.self_study)
        or (content.collaboration_status == "confirmed" and lp.collaboration)
        or (content.presentation_status == "confirmed" and lp.presentation)
    )
    if not has_any:
        return

    _add_section_title(document, "四、学习流程")

    if content.self_study_status == "confirmed" and lp.self_study:
        _render_assessment_items(document, "（一）自主学习（A级）", lp.self_study)

    if content.collaboration_status == "confirmed" and lp.collaboration:
        _render_assessment_items(document, "（二）合作探究（B级）", lp.collaboration)

    if content.presentation_status == "confirmed" and lp.presentation:
        _render_assessment_items(document, "（三）展示提升（C级）", lp.presentation)


def _render_assessment(document: DocxDocument, content: StudyGuideContent) -> None:
    if content.assessment_status != "confirmed" or not content.assessment:
        return
    _render_assessment_items(document, "五、达标测评", content.assessment)


def _render_extension(document: DocxDocument, content: StudyGuideContent) -> None:
    if content.extension_status != "confirmed" or not content.extension:
        return
    _render_assessment_items(document, "六、拓展延伸（D级）", content.extension)


def _render_self_reflection(document: DocxDocument, content: StudyGuideContent) -> None:
    _add_section_title(document, "七、自主反思")
    if content.self_reflection:
        _add_body(document, content.self_reflection)
    else:
        _add_empty_area(document, "（课后填写学习反思）")


def _build_study_guide_docx(
    document: DocxDocument,
    task: Task,
    content: StudyGuideContent,
) -> None:
    _render_study_guide_header(document, content, task)
    _render_learning_objectives(document, content)
    _render_key_difficulties(document, content)
    _render_prior_knowledge(document, content)
    _render_learning_process(document, content)
    _render_assessment(document, content)
    _render_extension(document, content)
    _render_self_reflection(document, content)


# ---------------------------------------------------------------------------
# 公共入口
# ---------------------------------------------------------------------------

def build_docx(task: Task, content: DocumentContent, template_spec: dict | None = None) -> bytes:
    """根据内容类型生成 Word 文档字节。"""
    document = DocxDocument()
    _configure_page(document)

    if is_lesson_plan(content):
        if template_spec and template_spec.get("kind") == "school_export_template":
            _build_templated_lesson_plan_docx(document, task, content, template_spec)
        else:
            _build_lesson_plan_docx(document, task, content)
    elif is_study_guide(content):
        _build_study_guide_docx(document, task, content)
    else:
        # 未知类型，写基本信息
        _add_paragraph(
            document,
            task.title,
            font_name="微软雅黑",
            font_size=_PT_EIGHTEEN,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
