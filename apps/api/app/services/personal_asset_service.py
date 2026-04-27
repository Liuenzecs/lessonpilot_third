from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from sqlmodel import Session, select

from app.models.personal_asset import PersonalAsset
from app.schemas.personal_asset import (
    ExtractedAssetSection,
    PersonalAssetConfirmPayload,
    PersonalAssetPreview,
    PersonalAssetRecommendation,
    ReuseSuggestion,
)

MAX_ASSET_BYTES = 10 * 1024 * 1024


def preview_personal_asset(file_bytes: bytes, filename: str) -> PersonalAssetPreview:
    lower_name = filename.lower()
    if len(file_bytes) > MAX_ASSET_BYTES:
        raise HTTPException(status_code=413, detail="资料超过 10MB，请拆分后再上传")
    if lower_name.endswith(".docx"):
        return _preview_docx(file_bytes, filename)
    if lower_name.endswith(".pptx"):
        return _preview_pptx(file_bytes, filename)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="当前仅支持 .docx 和 .pptx 资料")


def create_personal_asset(
    session: Session,
    user_id: str,
    payload: PersonalAssetConfirmPayload,
) -> PersonalAsset:
    preview = payload.preview
    asset = PersonalAsset(
        user_id=user_id,
        title=(payload.title or preview.title).strip() or "未命名资料",
        asset_type=payload.asset_type or preview.asset_type,
        source_filename=preview.source_filename,
        file_type=preview.file_type,
        subject=(payload.subject or preview.subject).strip() or "语文",
        grade=(payload.grade or preview.grade).strip(),
        topic=(payload.topic or preview.topic).strip(),
        extracted_content={
            "source_filename": preview.source_filename,
            "extracted_sections": [item.model_dump() for item in preview.extracted_sections],
            "unmapped_sections": [item.model_dump() for item in preview.unmapped_sections],
            "warnings": preview.warnings,
        },
        reuse_suggestions=[item.model_dump() for item in preview.reuse_suggestions],
    )
    session.add(asset)
    session.commit()
    session.refresh(asset)
    return asset


def list_personal_assets(session: Session, user_id: str) -> list[PersonalAsset]:
    statement = select(PersonalAsset).where(PersonalAsset.user_id == user_id).order_by(PersonalAsset.updated_at.desc())
    return session.exec(statement).all()


def recommend_personal_assets(
    session: Session,
    user_id: str,
    *,
    subject: str = "",
    grade: str = "",
    topic: str = "",
    keywords: str = "",
    asset_ids: list[str] | None = None,
    limit: int = 6,
) -> list[PersonalAssetRecommendation]:
    """Return teacher-owned snippets that can be reused for the current topic."""
    assets = _load_candidate_assets(session, user_id, asset_ids)
    terms = _query_terms(subject=subject, grade=grade, topic=topic, keywords=keywords)
    recommendations: list[PersonalAssetRecommendation] = []

    for asset in assets:
        for section in _asset_sections(asset):
            score, matched_terms = _score_asset_section(asset, section, terms, subject=subject, grade=grade)
            if asset_ids and score <= 0:
                score = 1
            if score <= 0:
                continue
            recommendations.append(
                PersonalAssetRecommendation(
                    asset_id=asset.id,
                    title=asset.title,
                    asset_type=asset.asset_type,
                    file_type=asset.file_type,
                    source_filename=asset.source_filename,
                    subject=asset.subject,
                    grade=asset.grade,
                    topic=asset.topic,
                    section_title=section.title or asset.title,
                    section_type=section.section_type,
                    content_snippet=_clip_text(section.content or asset.title, 260),
                    score=score,
                    matched_terms=matched_terms,
                )
            )

    recommendations.sort(key=lambda item: (item.score, item.title), reverse=True)
    return recommendations[:limit]


def validate_personal_asset_ids(session: Session, user_id: str, asset_ids: list[str]) -> None:
    if not asset_ids:
        return
    assets = _load_candidate_assets(session, user_id, asset_ids)
    found_ids = {asset.id for asset in assets}
    missing_ids = [asset_id for asset_id in asset_ids if asset_id not in found_ids]
    if missing_ids:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Personal asset not found")


def format_personal_asset_context(recommendations: list[PersonalAssetRecommendation]) -> str:
    if not recommendations:
        return ""
    lines = [
        "## 我的资料库参考",
        "",
        "以下资料来自当前老师的个人资料库，只能用于本次老师自己的备课。",
        "如引用，请在内容对应位置标注 [cite:资料ID]。",
        "",
    ]
    for item in recommendations:
        lines.append(f"[我的资料] ID: {item.asset_id}")
        lines.append(f"资料：{item.title}")
        lines.append(f"文件：{item.source_filename}（{item.file_type}）")
        lines.append(f"片段：{item.section_title}")
        lines.append(f"内容：{item.content_snippet}")
        lines.append("")
    return "\n".join(lines)


def get_personal_asset(session: Session, asset_id: str, user_id: str) -> PersonalAsset:
    asset = session.exec(
        select(PersonalAsset).where(PersonalAsset.id == asset_id, PersonalAsset.user_id == user_id)
    ).first()
    if asset is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Personal asset not found")
    return asset


def delete_personal_asset(session: Session, asset_id: str, user_id: str) -> None:
    asset = get_personal_asset(session, asset_id, user_id)
    session.delete(asset)
    session.commit()


def _load_candidate_assets(
    session: Session,
    user_id: str,
    asset_ids: list[str] | None = None,
) -> list[PersonalAsset]:
    statement = select(PersonalAsset).where(PersonalAsset.user_id == user_id)
    if asset_ids:
        statement = statement.where(PersonalAsset.id.in_(asset_ids))
    return list(session.exec(statement.order_by(PersonalAsset.updated_at.desc())).all())


def _asset_sections(asset: PersonalAsset) -> list[ExtractedAssetSection]:
    content = asset.extracted_content or {}
    raw_sections = list(content.get("extracted_sections") or [])
    raw_sections.extend(content.get("unmapped_sections") or [])
    sections: list[ExtractedAssetSection] = []
    for raw in raw_sections:
        if isinstance(raw, ExtractedAssetSection):
            sections.append(raw)
        elif isinstance(raw, dict):
            sections.append(ExtractedAssetSection.model_validate(raw))
    if sections:
        return sections
    return [
        ExtractedAssetSection(
            title=asset.title,
            content=" ".join(item for item in [asset.topic, asset.title] if item),
            section_type="unknown",
        )
    ]


def _score_asset_section(
    asset: PersonalAsset,
    section: ExtractedAssetSection,
    terms: list[str],
    *,
    subject: str,
    grade: str,
) -> tuple[int, list[str]]:
    haystack = " ".join(
        [
            asset.title,
            asset.topic,
            asset.subject,
            asset.grade,
            asset.source_filename,
            section.title,
            section.content,
            section.section_type,
        ]
    ).lower()
    matched_terms = [term for term in terms if term.lower() in haystack]
    has_topic_match = bool(asset.topic and topic_like(asset.topic, terms))
    if not matched_terms and not has_topic_match:
        return 0, []
    score = len(matched_terms) * 10
    if subject and asset.subject and subject == asset.subject:
        score += 3
    if grade and asset.grade and grade == asset.grade:
        score += 2
    if has_topic_match:
        score += 6
    if section.section_type in {"objectives", "teaching_process", "assessment", "ppt_slide"}:
        score += 1
    return score, matched_terms


def topic_like(value: str, terms: list[str]) -> bool:
    normalized = _normalize_term(value)
    return bool(normalized) and any(normalized == term or normalized in term or term in normalized for term in terms)


def _query_terms(*, subject: str, grade: str, topic: str, keywords: str) -> list[str]:
    _ = subject, grade
    raw = " ".join([topic, keywords])
    terms: list[str] = []
    for value in re.split(r"[\s，。、“”‘’；;：:、,.!！?？\-——（）()]+", raw):
        normalized = _normalize_term(value)
        if not normalized or normalized in {"语文", "年级", "课题", "教学", "教案", "学案"}:
            continue
        terms.append(normalized)
    quoted = re.findall(r"《([^》]+)》", raw)
    terms.extend(_normalize_term(item) for item in quoted if _normalize_term(item))
    compact_topic = _normalize_term(topic)
    if compact_topic:
        terms.append(compact_topic)
    unique_terms: list[str] = []
    for term in terms:
        if term and term not in unique_terms:
            unique_terms.append(term)
    return unique_terms


def _normalize_term(value: str) -> str:
    return value.replace("《", "").replace("》", "").strip().lower()


def _clip_text(value: str, limit: int) -> str:
    cleaned = _clean_text(value)
    return cleaned[:limit] + "..." if len(cleaned) > limit else cleaned


def _preview_docx(file_bytes: bytes, filename: str) -> PersonalAssetPreview:
    try:
        docx = DocxDocument(BytesIO(file_bytes))
    except (BadZipFile, ValueError, KeyError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Word 资料解析失败") from None

    paragraphs = [_clean_text(p.text) for p in docx.paragraphs if _clean_text(p.text)]
    table_texts = []
    for table in docx.tables:
        rows = [
            " | ".join(_clean_text(cell.text) for cell in row.cells if _clean_text(cell.text))
            for row in table.rows
        ]
        table_text = "\n".join(row for row in rows if row)
        if table_text:
            table_texts.append(table_text)
    sections = _sections_from_lines(paragraphs + table_texts)
    asset_type = _classify_docx(paragraphs + table_texts)
    title = _guess_title(paragraphs, filename)
    return PersonalAssetPreview(
        source_filename=filename,
        file_type="docx",
        title=title,
        subject=_extract_label_value(paragraphs, ["学科"]) or "语文",
        grade=_extract_label_value(paragraphs, ["年级", "授课年级"]),
        topic=_extract_label_value(paragraphs, ["课题", "教学内容"]) or title.replace("教案", ""),
        asset_type=asset_type,
        extracted_sections=sections[:12],
        unmapped_sections=[] if sections else [
            ExtractedAssetSection(title="未识别正文", content="\n".join(paragraphs[:20]))
        ],
        reuse_suggestions=_reuse_suggestions(asset_type),
        warnings=[] if sections else ["没有识别到清晰栏目，已保留正文供老师确认。"],
    )


def _preview_pptx(file_bytes: bytes, filename: str) -> PersonalAssetPreview:
    try:
        texts = _extract_pptx_slide_text(file_bytes)
    except (BadZipFile, KeyError, ET.ParseError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="PPT 资料解析失败") from None
    sections = [
        ExtractedAssetSection(title=f"第 {idx} 页", content=text, section_type="ppt_slide")
        for idx, text in enumerate(texts, start=1)
        if text
    ]
    title = sections[0].content.splitlines()[0][:80] if sections else filename.rsplit(".", 1)[0]
    return PersonalAssetPreview(
        source_filename=filename,
        file_type="pptx",
        title=title,
        asset_type="ppt_outline",
        extracted_sections=sections[:30],
        unmapped_sections=[],
        reuse_suggestions=_reuse_suggestions("ppt_outline"),
        warnings=[] if sections else ["未从 PPT 中提取到文字内容。"],
    )


def _extract_pptx_slide_text(file_bytes: bytes) -> list[str]:
    texts: list[str] = []
    with ZipFile(BytesIO(file_bytes)) as archive:
        slide_names = sorted(
            [name for name in archive.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)],
            key=lambda item: int(re.search(r"slide(\d+)\.xml", item).group(1)),  # type: ignore[union-attr]
        )
        for slide_name in slide_names:
            root = ET.fromstring(archive.read(slide_name))
            parts = [
                node.text.strip()
                for node in root.iter()
                if node.tag.endswith("}t") and node.text and node.text.strip()
            ]
            texts.append("\n".join(parts))
    return texts


def _sections_from_lines(lines: list[str]) -> list[ExtractedAssetSection]:
    sections: list[ExtractedAssetSection] = []
    current_title = ""
    current_body: list[str] = []
    for line in lines:
        section_type = _section_type(line)
        if section_type != "unknown":
            if current_title or current_body:
                sections.append(
                    ExtractedAssetSection(
                        title=current_title or "未命名栏目",
                        content="\n".join(current_body).strip(),
                        section_type=_section_type(current_title),
                    )
                )
            current_title = _strip_heading_marker(line)
            current_body = []
        else:
            current_body.append(line)
    if current_title or current_body:
        sections.append(
            ExtractedAssetSection(
                title=current_title or "未命名栏目",
                content="\n".join(current_body).strip(),
                section_type=_section_type(current_title),
            )
        )
    return [item for item in sections if item.title or item.content]


def _classify_docx(lines: list[str]) -> str:
    text = "\n".join(lines)
    if any(label in text for label in ["教学目标", "教学过程", "板书设计"]):
        return "lesson_plan"
    if any(label in text for label in ["学习目标", "自主学习", "达标测评"]):
        return "study_guide"
    if any(label in text for label in ["导入语", "问题链", "课堂小结", "易错点"]):
        return "teaching_note"
    return "reference_material"


def _reuse_suggestions(asset_type: str) -> list[ReuseSuggestion]:
    if asset_type == "lesson_plan":
        return [
            ReuseSuggestion(
                target="lesson_plan",
                label="作为旧教案导入",
                reason="包含正式教案栏目，可转为待确认教案。",
            ),
            ReuseSuggestion(
                target="school_template",
                label="识别学校格式",
                reason="若格式固定，可保存为个人导出模板。",
            ),
        ]
    if asset_type == "study_guide":
        return [ReuseSuggestion(target="study_guide", label="复用为学案草稿", reason="学生任务和测评可继续编辑。")]
    if asset_type == "ppt_outline":
        return [
            ReuseSuggestion(
                target="teaching_package",
                label="作为上课包参考",
                reason="可复用 slide 标题和课堂问题。",
            )
        ]
    return [ReuseSuggestion(target="knowledge", label="作为私有参考资料", reason="只在当前账号内用于备课参考。")]


def _section_type(value: str) -> str:
    if any(label in value for label in ["教学目标", "学习目标", "目标"]):
        return "objectives"
    if any(label in value for label in ["重难点", "重点", "难点"]):
        return "key_points"
    if any(label in value for label in ["教学过程", "教学流程", "教学环节"]):
        return "teaching_process"
    if any(label in value for label in ["达标测评", "当堂检测", "练习"]):
        return "assessment"
    if any(label in value for label in ["板书"]):
        return "board_design"
    return "unknown"


def _guess_title(lines: list[str], filename: str) -> str:
    for line in lines[:8]:
        if len(line) <= 80 and any(label in line for label in ["教案", "学案", "讲义", "教学设计", "《"]):
            return line
    return filename.rsplit(".", 1)[0]


def _extract_label_value(lines: list[str], labels: list[str]) -> str:
    for line in lines[:20]:
        for label in labels:
            match = re.search(rf"{re.escape(label)}\s*[:：]\s*([^；;，,\s]+)", line)
            if match:
                return match.group(1).strip()
    return ""


def _strip_heading_marker(value: str) -> str:
    return re.sub(r"^[一二三四五六七八九十\d]+[、.．]\s*", "", value).strip()


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()
