from __future__ import annotations

import re
from io import BytesIO
from zipfile import BadZipFile

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from sqlmodel import Session, select

from app.models.template import Template, TemplateSection
from app.schemas.template import (
    TemplateConfirmPayload,
    TemplateCreate,
    TemplateFieldMapping,
    TemplatePreview,
    TemplateTableLayout,
    TemplateUpdate,
)

MAX_TEMPLATE_BYTES = 5 * 1024 * 1024

_FIELD_ALIASES: dict[str, list[str]] = {
    "header.title": ["课题", "课名", "教学内容", "标题"],
    "header.subject": ["学科"],
    "header.grade": ["年级", "授课年级"],
    "header.classHour": ["课时", "第几课时"],
    "header.lessonCategory": ["课型"],
    "header.teacher": ["教师", "授课教师", "执教者"],
    "objectives": ["教学目标", "学习目标", "三维目标", "核心素养目标"],
    "key_points.keyPoints": ["教学重点", "重点"],
    "key_points.difficulties": ["教学难点", "难点"],
    "preparation": ["教学准备", "课前准备", "教具准备"],
    "teaching_process": ["教学过程", "教学流程", "教学环节"],
    "board_design": ["板书设计", "板书"],
    "reflection": ["教学反思", "课后反思", "教后反思"],
}

_SECTION_ORDER = [
    ("objectives", ["教学目标", "学习目标", "三维目标", "核心素养目标"]),
    ("key_points", ["教学重难点", "教学重点", "教学难点", "重难点"]),
    ("preparation", ["教学准备", "课前准备", "教具准备"]),
    ("teaching_process", ["教学过程", "教学流程", "教学活动", "教学环节"]),
    ("board_design", ["板书设计", "板书"]),
    ("reflection", ["教学反思", "课后反思", "教后反思"]),
]


def get_template(session: Session, template_id: str) -> Template | None:
    return session.get(Template, template_id)


def get_accessible_template(session: Session, template_id: str, user_id: str) -> Template | None:
    template = session.get(Template, template_id)
    if template is None:
        return None
    if template.is_public or template.user_id in (None, user_id):
        return template
    return None

def get_templates(
    session: Session,
    skip: int = 0,
    limit: int = 100,
    subject: str | None = None,
    grade: str | None = None,
    is_public: bool | None = None,
) -> list[Template]:
    statement = select(Template)
    if subject:
        statement = statement.where(Template.subject == subject)
    if grade:
        statement = statement.where(Template.grade == grade)
    if is_public is not None:
        statement = statement.where(Template.is_public == is_public)
    statement = statement.offset(skip).limit(limit)
    return session.exec(statement).all()

def create_template(session: Session, template_in: TemplateCreate, user_id: str | None = None) -> Template:
    db_template = Template.model_validate(template_in, update={"sections": None, "user_id": user_id})
    session.add(db_template)
    session.commit()
    session.refresh(db_template)

    if template_in.sections:
        for section in template_in.sections:
            db_section = TemplateSection(
                template_id=db_template.id,
                section_name=section.section_name,
                order=section.order,
                description=section.description,
                prompt_hints=section.prompt_hints,
                schema_constraints=section.schema_constraints,
            )
            session.add(db_section)
        session.commit()
    
    return db_template


def preview_school_template(file_bytes: bytes, filename: str) -> TemplatePreview:
    if not filename.lower().endswith(".docx"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="当前仅支持 .docx 学校模板")
    if len(file_bytes) > MAX_TEMPLATE_BYTES:
        raise HTTPException(status_code=413, detail="模板超过 5MB，请精简后再上传")
    try:
        docx = DocxDocument(BytesIO(file_bytes))
    except (BadZipFile, ValueError, KeyError):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Word 模板解析失败，请确认上传的是有效 .docx 文件",
        ) from None

    paragraphs = [_clean_text(paragraph.text) for paragraph in docx.paragraphs if _clean_text(paragraph.text)]
    rows = [[[_clean_text(cell.text) for cell in row.cells] for row in table.rows] for table in docx.tables]
    field_mappings = _detect_field_mappings(paragraphs, rows)
    section_order = _detect_section_order(paragraphs, rows)
    table_layouts = _detect_table_layouts(rows)
    blank_areas = _detect_blank_areas(paragraphs, rows)
    unsupported_items = _detect_unsupported_items(paragraphs)
    warnings: list[str] = []
    if "teaching_process" not in section_order:
        warnings.append("未明确识别到教学过程栏目，导出时会回退到默认教学过程位置。")
    if not table_layouts:
        warnings.append("未识别到可复用表格布局，模板导出将以栏目顺序为主。")

    title = _guess_template_name(paragraphs, filename)
    grade = _extract_label_value(paragraphs, ["年级", "授课年级"])
    subject = _extract_label_value(paragraphs, ["学科"]) or "语文"
    return TemplatePreview(
        source_filename=filename,
        name=title,
        subject=subject,
        grade=grade,
        template_type="lesson_plan",
        field_mappings=field_mappings,
        section_order=section_order,
        table_layouts=table_layouts,
        blank_areas=blank_areas,
        unsupported_items=unsupported_items,
        warnings=warnings,
    )


def save_school_template(session: Session, user_id: str, payload: TemplateConfirmPayload) -> Template:
    preview = payload.preview
    content = {
        "kind": "school_export_template",
        "spec_version": 1,
        "source_filename": preview.source_filename,
        "field_mappings": [item.model_dump() for item in preview.field_mappings],
        "section_order": preview.section_order,
        "table_layouts": [item.model_dump() for item in preview.table_layouts],
        "blank_areas": preview.blank_areas,
        "unsupported_items": preview.unsupported_items,
        "warnings": preview.warnings,
    }
    template = Template(
        user_id=user_id,
        name=(payload.name or preview.name).strip() or "学校教案模板",
        subject=(payload.subject or preview.subject).strip() or "语文",
        grade=(payload.grade or preview.grade).strip(),
        description="由学校 Word 模板识别保存",
        template_type="school_lesson_export",
        content=content,
        is_public=False,
    )
    session.add(template)
    session.commit()
    session.refresh(template)
    return template


def list_school_templates(session: Session, user_id: str) -> list[Template]:
    statement = (
        select(Template)
        .where(Template.user_id == user_id, Template.template_type == "school_lesson_export")
        .order_by(Template.updated_at.desc())
    )
    return session.exec(statement).all()


def delete_school_template(session: Session, template_id: str, user_id: str) -> None:
    template = session.exec(
        select(Template).where(
            Template.id == template_id,
            Template.user_id == user_id,
            Template.template_type == "school_lesson_export",
        )
    ).first()
    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
    session.delete(template)
    session.commit()

def update_template(session: Session, db_template: Template, template_in: TemplateUpdate) -> Template:
    update_data = template_in.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_template, field, value)
    
    session.add(db_template)
    session.commit()
    session.refresh(db_template)
    return db_template

def get_template_sections(session: Session, template_id: str) -> list[TemplateSection]:
    statement = (
        select(TemplateSection)
        .where(TemplateSection.template_id == template_id)
        .order_by(TemplateSection.order)
    )
    return session.exec(statement).all()


def _detect_field_mappings(
    paragraphs: list[str],
    tables: list[list[list[str]]],
) -> list[TemplateFieldMapping]:
    mappings: list[TemplateFieldMapping] = []
    seen: set[tuple[str, str]] = set()
    candidates = paragraphs[:40] + [cell for table in tables for row in table[:8] for cell in row]
    for text in candidates:
        for field, aliases in _FIELD_ALIASES.items():
            for alias in aliases:
                if alias in text and (alias, field) not in seen:
                    seen.add((alias, field))
                    mappings.append(
                        TemplateFieldMapping(
                            template_label=alias,
                            content_field=field,
                            confidence=0.9,
                    location=_mapping_location(alias, tables),
                        )
                    )
    return mappings


def _detect_section_order(paragraphs: list[str], tables: list[list[list[str]]]) -> list[str]:
    order: list[str] = []
    candidates = paragraphs + [cell for table in tables for row in table for cell in row]
    for text in candidates:
        normalized = _normalize_heading(text)
        for section, aliases in _SECTION_ORDER:
            if section in order:
                continue
            if any(_normalize_heading(alias) in normalized for alias in aliases):
                order.append(section)
    return order


def _mapping_location(alias: str, tables: list[list[list[str]]]) -> str:
    if any(alias in cell for table in tables for row in table for cell in row):
        return "table"
    return "paragraph"


def _detect_table_layouts(tables: list[list[list[str]]]) -> list[TemplateTableLayout]:
    layouts: list[TemplateTableLayout] = []
    for table in tables:
        rows = [row for row in table if any(row)]
        if not rows:
            continue
        header_text = " ".join(" ".join(row) for row in rows[:2])
        if any(word in header_text for word in ["教师活动", "学生活动", "设计意图", "教学环节"]):
            columns = [cell for cell in rows[0] if cell] or _best_effort_columns(rows)
            layouts.append(
                TemplateTableLayout(
                    name="teaching_process",
                    columns=columns,
                    field_mappings=[
                        TemplateFieldMapping(
                            template_label=col,
                            content_field=_process_column_field(col),
                            location="table",
                        )
                        for col in columns
                    ],
                )
            )
        elif any(word in header_text for word in ["学科", "年级", "课题", "课时", "教师"]):
            labels = [cell for row in rows for cell in row if _looks_like_field_label(cell)]
            layouts.append(
                TemplateTableLayout(
                    name="metadata",
                    columns=labels,
                    field_mappings=[
                        TemplateFieldMapping(
                            template_label=label,
                            content_field=_metadata_field(label),
                            location="table",
                        )
                        for label in labels
                    ],
                )
            )
    return layouts


def _detect_blank_areas(paragraphs: list[str], tables: list[list[list[str]]]) -> list[str]:
    text_items = paragraphs + [cell for table in tables for row in table for cell in row]
    blank_labels = []
    for text in text_items:
        for label in ["教学反思", "教后反思", "二次备课", "听课评价", "签字", "审核"]:
            if label in text and label not in blank_labels:
                blank_labels.append(text.strip("：: ") if len(text) <= 40 else label)
    return blank_labels


def _detect_unsupported_items(paragraphs: list[str]) -> list[str]:
    unsupported = []
    for text in paragraphs:
        if any(label in text for label in ["作业设计", "二次备课", "审核意见", "签字"]):
            unsupported.append(text[:120])
    return unsupported[:10]


def _guess_template_name(paragraphs: list[str], filename: str) -> str:
    for text in paragraphs[:8]:
        if "教案" in text or "教学设计" in text or "模板" in text:
            return text[:80]
    return filename.rsplit(".", 1)[0]


def _extract_label_value(paragraphs: list[str], labels: list[str]) -> str:
    for text in paragraphs[:20]:
        for label in labels:
            match = re.search(rf"{re.escape(label)}\s*[:：]\s*([^；;，,\s]+)", text)
            if match:
                return match.group(1).strip()
    return ""


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()


def _normalize_heading(value: str) -> str:
    value = re.sub(r"^[一二三四五六七八九十\d]+[、.．]\s*", "", value.strip())
    return re.sub(r"[:：\s]", "", value)


def _looks_like_field_label(value: str) -> bool:
    return any(label in value for label in ["学科", "年级", "课题", "课时", "课型", "教师"])


def _metadata_field(label: str) -> str:
    for field, aliases in _FIELD_ALIASES.items():
        if field.startswith("header") and any(alias in label for alias in aliases):
            return field
    return "unsupported"


def _process_column_field(label: str) -> str:
    if "环节" in label or "流程" in label or "步骤" in label:
        return "teaching_process.phase"
    if "时" in label or "分钟" in label:
        return "teaching_process.duration"
    if "教师" in label or "教法" in label:
        return "teaching_process.teacher_activity"
    if "学生" in label or "学法" in label:
        return "teaching_process.student_activity"
    if "意图" in label or "说明" in label:
        return "teaching_process.design_intent"
    return "teaching_process.extra"


def _best_effort_columns(rows: list[list[str]]) -> list[str]:
    return [cell for row in rows[:2] for cell in row if cell][:6]
