from __future__ import annotations

import re
from io import BytesIO
from zipfile import BadZipFile

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from sqlmodel import Session

from app.models import Document, Task
from app.schemas.content import (
    KeyPoints,
    LessonPlanContent,
    LessonPlanHeader,
    TeachingObjective,
    TeachingProcessStep,
)
from app.schemas.document import DocumentRead
from app.schemas.lesson_import import (
    ImportWarning,
    LessonPlanImportConfirmPayload,
    LessonPlanImportMetadata,
    LessonPlanImportPreview,
    UnmappedSection,
)
from app.schemas.task import TaskRead
from app.services.document_service import serialize_document

MAX_IMPORT_BYTES = 5 * 1024 * 1024

SECTION_ALIASES: dict[str, list[str]] = {
    "objectives": ["教学目标", "学习目标", "三维目标", "核心素养目标", "素养目标"],
    "key_points": ["教学重难点", "教学重点难点", "教学重点", "教学难点", "重难点"],
    "preparation": ["教学准备", "课前准备", "教具准备", "教学资源"],
    "teaching_process": ["教学过程", "教学流程", "课堂教学过程", "教学活动", "教学环节"],
    "board_design": ["板书设计", "板书"],
    "reflection": ["教学反思", "课后反思"],
}

SECTION_TITLES = {
    "objectives": "教学目标",
    "key_points": "教学重难点",
    "preparation": "教学准备",
    "teaching_process": "教学过程",
    "board_design": "板书设计",
    "reflection": "教学反思",
}


def preview_lesson_plan_import(file_bytes: bytes, filename: str) -> LessonPlanImportPreview:
    if not filename.lower().endswith(".docx"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="当前仅支持 .docx 教案导入")
    if len(file_bytes) > MAX_IMPORT_BYTES:
        raise HTTPException(status_code=413, detail="文件超过 5MB，请拆分后再导入")

    try:
        docx = DocxDocument(BytesIO(file_bytes))
    except (BadZipFile, ValueError, KeyError):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Word 文件解析失败，请确认上传的是有效 .docx 文件",
        ) from None

    paragraphs = [_clean_text(p.text) for p in docx.paragraphs if _clean_text(p.text)]
    metadata = _extract_metadata(paragraphs, filename)
    buckets, unmapped, warnings = _bucket_document_items(docx, paragraphs)
    content = _build_lesson_plan_content(metadata, buckets, warnings)
    mapped_sections = _mapped_sections(content)

    if not mapped_sections:
        warnings.append(ImportWarning(message="没有识别到标准教案栏目，已把正文放入未识别内容，请手动检查。"))

    for section in ("objectives", "key_points", "teaching_process"):
        if section not in mapped_sections:
            warnings.append(
                ImportWarning(section=section, message=f"未识别到{SECTION_TITLES[section]}，导入后可能需要补写。")
            )

    return LessonPlanImportPreview(
        source_filename=filename,
        metadata=metadata,
        content=content,
        mapped_sections=mapped_sections,
        unmapped_sections=unmapped,
        warnings=warnings,
    )


def create_imported_lesson_plan(
    session: Session,
    user_id: str,
    payload: LessonPlanImportConfirmPayload,
) -> tuple[TaskRead, DocumentRead]:
    metadata = payload.metadata
    content = payload.content
    title = (metadata.title or metadata.topic or content.header.title or "导入教案").strip()
    topic = (metadata.topic or content.header.title or title).strip()

    task = Task(
        user_id=user_id,
        title=title,
        subject=metadata.subject.strip() or "语文",
        grade=metadata.grade.strip(),
        topic=topic,
        requirements="由旧 Word 教案导入",
        status="imported",
        scene=metadata.scene,
        lesson_type="lesson_plan",
        class_hour=metadata.class_hour,
        lesson_category=metadata.lesson_category,
    )
    session.add(task)
    session.flush()

    content.header = LessonPlanHeader(
        title=title,
        subject=task.subject,
        grade=task.grade,
        class_hour=task.class_hour,
        lesson_category=task.lesson_category,  # type: ignore[arg-type]
        teacher=content.header.teacher,
    )
    _force_pending(content)

    document = Document(
        task_id=task.id,
        user_id=user_id,
        doc_type="lesson_plan",
        title=title,
        content=content.model_dump(by_alias=True),
    )
    session.add(document)
    session.commit()
    session.refresh(task)
    session.refresh(document)

    task_read = TaskRead(
        id=task.id,
        title=task.title,
        subject=task.subject,
        grade=task.grade,
        topic=task.topic,
        requirements=task.requirements,
        status=task.status,
        scene=task.scene,
        lesson_type=task.lesson_type,
        class_hour=task.class_hour,
        lesson_category=task.lesson_category,
        template_id=task.template_id,
        created_at=task.created_at,
        updated_at=task.updated_at,
    )
    return task_read, serialize_document(document)


def _bucket_document_items(
    docx,
    paragraphs: list[str],
) -> tuple[dict[str, list[str | list[list[str]]]], list[UnmappedSection], list[ImportWarning]]:
    buckets: dict[str, list[str | list[list[str]]]] = {key: [] for key in SECTION_ALIASES}
    unmapped: list[UnmappedSection] = []
    warnings: list[ImportWarning] = []
    current_section: str | None = None
    consumed_paragraphs: set[str] = set()

    for text in paragraphs:
        section, inline = _detect_section_heading(text)
        if section:
            current_section = section
            consumed_paragraphs.add(text)
            if inline:
                buckets[section].append(inline)
            continue
        if current_section:
            buckets[current_section].append(text)
            consumed_paragraphs.add(text)
            continue
        if _looks_like_metadata(text):
            consumed_paragraphs.add(text)

    for table in docx.tables:
        rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = [[cell for cell in row] for row in rows if any(row)]
        if not rows:
            continue
        if _table_looks_like_process(rows):
            buckets["teaching_process"].append(rows)
        elif current_section:
            buckets[current_section].append(_table_to_text(rows))
            warnings.append(
                ImportWarning(section=current_section, message="有表格未能识别为教学过程，已作为普通文本导入。")
            )
        else:
            unmapped.append(UnmappedSection(title="未识别表格", content=_table_to_text(rows)))

    for text in paragraphs:
        if text not in consumed_paragraphs and not _looks_like_title(text):
            unmapped.append(UnmappedSection(title=None, content=text))

    if unmapped:
        warnings.append(ImportWarning(message=f"有 {len(unmapped)} 段内容未能确定归属，请在预览中检查。"))
    return buckets, unmapped, warnings


def _build_lesson_plan_content(
    metadata: LessonPlanImportMetadata,
    buckets: dict[str, list[str | list[list[str]]]],
    warnings: list[ImportWarning],
) -> LessonPlanContent:
    content = LessonPlanContent(
        header=LessonPlanHeader(
            title=metadata.title,
            subject=metadata.subject,
            grade=metadata.grade,
            class_hour=metadata.class_hour,
            lesson_category=metadata.lesson_category,
        )
    )
    content.objectives = _parse_objectives(_string_items(buckets["objectives"]))
    content.key_points = _parse_key_points(_string_items(buckets["key_points"]))
    content.preparation = _split_lines(_string_items(buckets["preparation"]))
    content.teaching_process = _parse_teaching_process(buckets["teaching_process"], warnings)
    content.board_design = "\n".join(_split_lines(_string_items(buckets["board_design"]))).strip()
    content.reflection = "\n".join(_split_lines(_string_items(buckets["reflection"]))).strip()
    _force_pending(content)
    return content


def _parse_objectives(items: list[str]) -> list[TeachingObjective]:
    lines = _split_lines(items)
    dimensions = ["knowledge", "ability", "emotion"]
    objectives: list[TeachingObjective] = []
    for index, line in enumerate(lines[:6]):
        lower = line.lower()
        if "情感" in line or "价值观" in line or "emotion" in lower:
            dimension = "emotion"
        elif "过程" in line or "方法" in line or "能力" in line or "ability" in lower:
            dimension = "ability"
        else:
            dimension = dimensions[min(index, len(dimensions) - 1)]
        objectives.append(TeachingObjective(dimension=dimension, content=_strip_list_marker(line)))
    return objectives


def _parse_key_points(items: list[str]) -> KeyPoints:
    lines = _split_lines(items)
    key_points: list[str] = []
    difficulties: list[str] = []
    current = "key"
    for line in lines:
        clean = _strip_list_marker(line)
        if "难点" in clean:
            current = "difficulty"
            value = _strip_label(clean, ["教学难点", "难点"])
        elif "重点" in clean:
            current = "key"
            value = _strip_label(clean, ["教学重点", "重点"])
        else:
            value = clean
        if not value:
            continue
        if current == "difficulty":
            difficulties.append(value)
        else:
            key_points.append(value)
    return KeyPoints(key_points=key_points, difficulties=difficulties)


def _parse_teaching_process(
    items: list[str | list[list[str]]],
    warnings: list[ImportWarning],
) -> list[TeachingProcessStep]:
    steps: list[TeachingProcessStep] = []
    text_lines: list[str] = []
    for item in items:
        if isinstance(item, list):
            parsed = _parse_process_table(item, warnings)
            steps.extend(parsed)
        else:
            text_lines.extend(_split_lines([item]))

    for line in text_lines:
        clean = _strip_list_marker(line)
        if not clean:
            continue
        steps.append(
            TeachingProcessStep(
                phase=clean[:40],
                duration=_parse_duration(clean) or 5,
                teacher_activity=clean,
                student_activity="",
                design_intent="",
                status="pending",
            )
        )
    if text_lines and not any(step.student_activity for step in steps):
        warnings.append(
            ImportWarning(section="teaching_process", message="段落式教学过程缺少学生活动列，已保守导入为过程草稿。")
        )
    return steps


def _parse_process_table(rows: list[list[str]], warnings: list[ImportWarning]) -> list[TeachingProcessStep]:
    header = rows[0]
    mapping = {
        "phase": _find_column(header, ["环节", "流程", "步骤", "阶段"]),
        "duration": _find_column(header, ["时间", "时长", "分钟"]),
        "teacher_activity": _find_column(header, ["教师活动", "教师", "教法", "师生活动"]),
        "student_activity": _find_column(header, ["学生活动", "学生", "学法", "师生活动"]),
        "design_intent": _find_column(header, ["设计意图", "意图", "设计说明"]),
    }
    data_rows = rows[1:] if any(value is not None for value in mapping.values()) else rows
    if mapping["phase"] is None:
        mapping["phase"] = 0
    if mapping["duration"] is None and len(header) > 1:
        mapping["duration"] = 1
    if mapping["teacher_activity"] is None and len(header) > 2:
        mapping["teacher_activity"] = 2
    if mapping["student_activity"] is None and len(header) > 3:
        mapping["student_activity"] = 3
    if mapping["design_intent"] is None and len(header) > 4:
        mapping["design_intent"] = 4

    steps: list[TeachingProcessStep] = []
    for index, row in enumerate(data_rows, start=1):
        phase = _cell_at(row, mapping["phase"]) or f"环节{index}"
        teacher = _cell_at(row, mapping["teacher_activity"])
        student = _cell_at(row, mapping["student_activity"])
        design = _cell_at(row, mapping["design_intent"])
        duration = _parse_duration(_cell_at(row, mapping["duration"])) or _parse_duration(" ".join(row)) or 5
        if not teacher and len(row) == 1:
            teacher = row[0]
        steps.append(
            TeachingProcessStep(
                phase=phase,
                duration=duration,
                teacher_activity=teacher,
                student_activity=student,
                design_intent=design,
                status="pending",
            )
        )
    if steps and any(not step.student_activity for step in steps):
        warnings.append(ImportWarning(section="teaching_process", message="部分教学过程行缺少学生活动，请导入后检查。"))
    if steps and any(not step.design_intent for step in steps):
        warnings.append(
            ImportWarning(section="teaching_process", message="部分教学过程行缺少设计意图，公立校提交前建议补齐。")
        )
    return steps


def _extract_metadata(paragraphs: list[str], filename: str) -> LessonPlanImportMetadata:
    title = ""
    subject = "语文"
    grade = ""
    topic = ""
    class_hour = 1
    lesson_category = "new"

    for text in paragraphs[:20]:
        title = title or _extract_value(text, ["标题", "课题名称"])
        topic = topic or _extract_value(text, ["课题", "教学内容"])
        subject = _extract_value(text, ["学科"]) or subject
        grade = _extract_value(text, ["年级", "授课年级"]) or grade
        hour_match = re.search(r"([一二三四五六七八九十\d]+)\s*课时|课时\s*[:：]?\s*(\d+)", text)
        if hour_match:
            class_hour = _chinese_number_to_int(hour_match.group(1) or hour_match.group(2))
        if "复习" in text:
            lesson_category = "review"
        elif "练习" in text:
            lesson_category = "exercise"
        elif "综合" in text:
            lesson_category = "comprehensive"

    title = title or topic or _first_title(paragraphs) or filename.rsplit(".", 1)[0]
    topic = topic or title.replace("教案", "").strip("《》 ")
    return LessonPlanImportMetadata(
        title=title[:255],
        subject=subject[:80],
        grade=grade[:80],
        topic=topic[:255],
        class_hour=max(1, min(class_hour, 10)),
        lesson_category=lesson_category,  # type: ignore[arg-type]
    )


def _detect_section_heading(text: str) -> tuple[str | None, str]:
    normalized = _normalize_heading(text)
    for section, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            alias_norm = _normalize_heading(alias)
            if normalized == alias_norm:
                return section, ""
            if normalized.startswith(alias_norm):
                inline = re.sub(rf"^\s*[一二三四五六七八九十\d、.．\s]*{re.escape(alias)}\s*[:：]?", "", text).strip()
                return section, inline
    return None, ""


def _mapped_sections(content: LessonPlanContent) -> list[str]:
    mapped: list[str] = []
    if content.objectives:
        mapped.append("objectives")
    if content.key_points.key_points or content.key_points.difficulties:
        mapped.append("key_points")
    if content.preparation:
        mapped.append("preparation")
    if content.teaching_process:
        mapped.append("teaching_process")
    if content.board_design:
        mapped.append("board_design")
    if content.reflection:
        mapped.append("reflection")
    return mapped


def _force_pending(content: LessonPlanContent) -> None:
    content.objectives_status = "pending"
    content.key_points_status = "pending"
    content.preparation_status = "pending"
    content.teaching_process_status = "pending"
    content.board_design_status = "pending"
    content.reflection_status = "pending"
    for step in content.teaching_process:
        step.status = "pending"


def _string_items(items: list[str | list[list[str]]]) -> list[str]:
    return [item for item in items if isinstance(item, str)]


def _split_lines(items: list[str]) -> list[str]:
    lines: list[str] = []
    for item in items:
        for part in re.split(r"[\n\r]+", item):
            clean = _clean_text(part)
            if clean:
                lines.append(clean)
    return lines


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()


def _normalize_heading(value: str) -> str:
    value = re.sub(r"^[一二三四五六七八九十\d]+[、.．]\s*", "", value.strip())
    value = re.sub(r"[:：\s]", "", value)
    return value


def _strip_list_marker(value: str) -> str:
    return re.sub(r"^[（(]?[一二三四五六七八九十\d]+[）)、.．]\s*", "", value).strip()


def _strip_label(value: str, labels: list[str]) -> str:
    for label in labels:
        value = re.sub(rf"^{re.escape(label)}\s*[:：]?", "", value).strip()
    return value


def _extract_value(text: str, labels: list[str]) -> str:
    for label in labels:
        match = re.search(rf"{re.escape(label)}\s*[:：]\s*([^；;，,\s]+)", text)
        if match:
            return match.group(1).strip()
    return ""


def _first_title(paragraphs: list[str]) -> str:
    for text in paragraphs[:8]:
        if _looks_like_title(text):
            return text
    return paragraphs[0] if paragraphs else ""


def _looks_like_title(text: str) -> bool:
    if len(text) > 80:
        return False
    return "教案" in text or "教学设计" in text or text.startswith("《")


def _looks_like_metadata(text: str) -> bool:
    return any(label in text for label in ["学科", "年级", "课题", "课时", "授课教师", "教师"])


def _table_looks_like_process(rows: list[list[str]]) -> bool:
    first_two = " ".join(" ".join(row) for row in rows[:2])
    return any(word in first_two for word in ["教师活动", "学生活动", "设计意图", "教学环节", "教学过程"])


def _table_to_text(rows: list[list[str]]) -> str:
    return "\n".join(" | ".join(cell for cell in row if cell) for row in rows)


def _find_column(header: list[str], keywords: list[str]) -> int | None:
    for index, cell in enumerate(header):
        if any(keyword in cell for keyword in keywords):
            return index
    return None


def _cell_at(row: list[str], index: int | None) -> str:
    if index is None or index >= len(row):
        return ""
    return row[index]


def _parse_duration(value: str) -> int | None:
    match = re.search(r"(\d+)", value or "")
    if not match:
        return None
    return max(1, min(int(match.group(1)), 180))


def _chinese_number_to_int(value: str | None) -> int:
    if not value:
        return 1
    if value.isdigit():
        return int(value)
    mapping = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
    if value == "十":
        return 10
    if value.startswith("十"):
        return 10 + mapping.get(value[-1], 0)
    if "十" in value:
        left, right = value.split("十", 1)
        return mapping.get(left, 1) * 10 + mapping.get(right, 0)
    return mapping.get(value, 1)
